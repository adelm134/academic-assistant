import streamlit as st
import requests
import json
import os
import io
import re
import wave
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches
import pypandoc
import openai

from vosk import Model, KaldiRecognizer
from streamlit_webrtc import webrtc_streamer, AudioProcessorBase

# -------------------------------
# Конфигурация пользователей для авторизации
VALID_USERS = {
    "user": "password",  # логин: user, пароль: password
}

# Файл для хранения истории чатов
CHAT_HISTORY_FILE = "chat_history.json"

# -------------------------------
# Функции работы с историей чатов
def load_chat_history():
    if os.path.exists(CHAT_HISTORY_FILE):
        with open(CHAT_HISTORY_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_chat_history(entry):
    history = load_chat_history()
    history.append(entry)
    with open(CHAT_HISTORY_FILE, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=4)

# -------------------------------
# Авторизация
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    st.sidebar.header("Login")
    username = st.sidebar.text_input("Username")
    password = st.sidebar.text_input("Password", type="password")
    if st.sidebar.button("Login"):
        if username in VALID_USERS and VALID_USERS[username] == password:
            st.session_state.authenticated = True
            st.session_state.username = username
            st.sidebar.success("Logged in as " + username)
        else:
            st.sidebar.error("Invalid username or password")
    st.stop()

# -------------------------------
# Выбор языка: английский или русский
language = st.selectbox("Choose language:", ("English", "Russian"))

# Определяем путь к модели Vosk в зависимости от выбранного языка
if language == "English":
    MODEL_PATH = "D:/voskvosk/model/vosk-model-small-en-us-0.15"
else:
    MODEL_PATH = r"D:/voskvosk/model/vosk-model-small-ru-0.22/vosk-model-small-ru-0.22"

# -------------------------------
# Инициализация Vosk модели
@st.cache_resource
def load_vosk_model(model_path):
    if not os.path.exists(model_path):
        st.error("Vosk model not found at the specified path!")
        return None
    return Model(model_path)

model = load_vosk_model(MODEL_PATH)

# -------------------------------
# Функция распознавания речи из аудиофайла (для тестирования)
def transcribe_audio(audio_bytes):
    wf = wave.open(io.BytesIO(audio_bytes), "rb")
    rec = KaldiRecognizer(model, wf.getframerate())
    result_text = ""
    while True:
        data = wf.readframes(4000)
        if len(data) == 0:
            break
        if rec.AcceptWaveform(data):
            res = json.loads(rec.Result())
            result_text += " " + res.get("text", "")
    res = json.loads(rec.FinalResult())
    result_text += " " + res.get("text", "")
    return result_text.strip()

# -------------------------------
# Класс для обработки аудио с микрофона в реальном времени
class VoskAudioProcessor(AudioProcessorBase):
    def __init__(self):
        self.model = load_vosk_model(MODEL_PATH)
        self.rec = KaldiRecognizer(self.model, 16000)
        self.transcript = ""
    def recv(self, frame):
        audio = frame.to_ndarray().tobytes()
        if self.rec.AcceptWaveform(audio):
            res = json.loads(self.rec.Result())
            self.transcript += " " + res.get("text", "")
        return frame

# -------------------------------
# Функция для генерации текста через DeepSeek API с использованием OpenAI SDK
def generate_text(prompt):
    openai.api_key = "sk-a4c644b2b7f349bb8a72a5cf9beebfe6"
    openai.api_base = "https://api.deepseek.com"
    if language == "Russian":
        system_msg = ("Вы являетесь помощником, генерирующим подробные и структурированные научные статьи "
                      "в формате IMRAD с гуманизированным тоном. Текст должен быть цельным, без использования списков. "
                      "Включите цитирования в квадратных скобках и литературный обзор с 10 источниками.")
    else:
        system_msg = ("You are a helpful assistant that generates detailed and well-structured academic texts in IMRAD format "
                      "with a humanized tone. The text should be continuous without bullet lists. Include citations in square brackets "
                      "and a literature overview with 10 references.")
    try:
        response = openai.ChatCompletion.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": prompt},
            ],
            stream=False
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error("Text generation failed: " + str(e))
        return ""

# -------------------------------
# Функция для генерации обзора для найденного источника
def generate_reference_overview(title, journal, doi):
    if language == "Russian":
        return f"В источнике '{title}', опубликованном в журнале {journal}, представлен анализ по данной тематике. DOI: {doi}."
    else:
        return f"The source '{title}', published in {journal}, provides an analysis on the topic. DOI: {doi}."

# -------------------------------
# Функция для поиска литературы через CrossRef API (10 источников)
def search_literature(query, rows=10):
    url = "https://api.crossref.org/works"
    params = {"query": query, "rows": rows}
    response = requests.get(url, params=params)
    results = []
    if response.status_code == 200:
        items = response.json()['message']['items']
        for idx, item in enumerate(items, start=1):
            title = item.get("title", ["No title"])[0]
            doi = item.get("DOI", "N/A")
            journal = item.get("container-title", [""])[0]
            overview = generate_reference_overview(title, journal, doi)
            results.append({"index": idx, "title": title, "doi": doi, "journal": journal, "overview": overview})
    else:
        st.error("Literature search failed: " + response.text)
    return results

# -------------------------------
# Функция для извлечения данных для визуализации из сгенерированной статьи
def extract_visualization_data(text):
    pattern = r"```json(.*?)```"
    match = re.search(pattern, text, re.DOTALL)
    if match:
        json_str = match.group(1).strip()
        try:
            data = json.loads(json_str)
            if "data_for_visualization" in data:
                return data["data_for_visualization"]
            else:
                st.warning("JSON block found but key 'data_for_visualization' is missing.")
                return None
        except Exception as e:
            st.error("Failed to parse visualization data JSON: " + str(e))
            return None
    else:
        st.info("No visualization data found in the generated article.")
        return None

# -------------------------------
# Функция для построения визуализаций: столбчатой диаграммы и линейного графика
def plot_generated_data(data):
    try:
        df = pd.DataFrame(data)
        st.write("### Data used for visualizations:")
        st.dataframe(df)
        # Столбчатая диаграмма
        plt.figure(figsize=(6,4))
        if "Category" in df.columns and "Value" in df.columns:
            sns.barplot(x="Category", y="Value", data=df)
            bar_chart_path = "bar_chart.png"
            plt.savefig(bar_chart_path)
            st.image(bar_chart_path, caption="Bar Chart")
            plt.clf()
        else:
            st.warning("Data does not contain 'Category' and 'Value' columns for bar plot.")

        # Линейный график
        plt.figure(figsize=(6,4))
        if "Category" in df.columns and "Value" in df.columns:
            sns.lineplot(x="Category", y="Value", data=df, marker="o")
            line_chart_path = "line_chart.png"
            plt.savefig(line_chart_path)
            st.image(line_chart_path, caption="Line Chart")
            plt.clf()
        else:
            st.warning("Data does not contain 'Category' and 'Value' columns for line plot.")
        return bar_chart_path, line_chart_path
    except Exception as e:
        st.error("Visualization failed: " + str(e))
        return None, None

# -------------------------------
# Функция для экспорта работы в DOCX с добавлением таблиц, графиков и литературного обзора
def export_document(article_text, vis_data, references, filename="article"):
    doc = Document()
    doc.add_heading("Scientific Article", level=0)
    # Добавляем сгенерированный текст статьи как сплошной текст
    doc.add_paragraph(article_text)
    
    if vis_data:
        doc.add_heading("Data Table", level=1)
        df = pd.DataFrame(vis_data)
        if not df.empty:
            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = str(col)
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, col in enumerate(df.columns):
                    row_cells[i].text = str(row[col])
    
    bar_chart_path, line_chart_path = plot_generated_data(vis_data) if vis_data else (None, None)
    doc.add_heading("Visualizations", level=1)
    if bar_chart_path and os.path.exists(bar_chart_path):
        doc.add_paragraph("Bar Chart:")
        doc.add_picture(bar_chart_path, width=Inches(5))
    if line_chart_path and os.path.exists(line_chart_path):
        doc.add_paragraph("Line Chart:")
        doc.add_picture(line_chart_path, width=Inches(5))
    
    if references:
        doc.add_heading("Literature Overview", level=1)
        for ref in references:
            ref_text = f"{ref['overview']}"
            doc.add_paragraph(ref_text)
    
    docx_filename = f"{filename}.docx"
    doc.save(docx_filename)
    return docx_filename

# -------------------------------
# Боковая панель: история чатов с возможностью раскрытия полной информации
st.sidebar.header("Chat History")
history = load_chat_history()
if history:
    for entry in reversed(history):
        with st.sidebar.expander(f"{entry['timestamp']} - {entry['username']}"):
            st.markdown("**Input:**")
            st.write(entry["input"])
            st.markdown("**Generated Article:**")
            st.write(entry["generated"])
else:
    st.sidebar.info("No chat history found.")

# -------------------------------
# Основной интерфейс
st.title("Scientific Article Generator")

# Выбор способа ввода: текст или голос (реальное время)
input_mode = st.radio("Choose input mode:", ("Text Input", "Voice Input (Real-Time)"))

article_text = ""
references = []
submitted = False

if input_mode == "Text Input":
    st.header("Enter your article details:")
    with st.form("article_questions"):
        st.write("Please answer the following questions in a continuous text:")
        if language == "Russian":
            q1 = st.text_input("1. Какую проблему вы исследуете?")
            q2 = st.text_input("2. Почему это важно?")
            q3 = st.text_input("3. Какие исследования уже проводились по данной теме?")
            q4 = st.text_input("4. Какие методы вы использовали?")
            q5 = st.text_input("5. Как вы собирали данные?")
            q6 = st.text_input("6. Какие результаты вы получили?")
            q7 = st.text_input("7. Соответствуют ли они ожиданиям или другим исследованиям?")
            q8 = st.text_input("8. Какие выводы можно сделать?")
            q9 = st.text_input("9. Какие дальнейшие исследования можно провести?")
        else:
            q1 = st.text_input("1. What problem are you researching?")
            q2 = st.text_input("2. Why is it important?")
            q3 = st.text_input("3. What previous research exists on this topic?")
            q4 = st.text_input("4. What methods did you use?")
            q5 = st.text_input("5. How did you collect data?")
            q6 = st.text_input("6. What results did you get?")
            q7 = st.text_input("7. Do they match expectations or other studies?")
            q8 = st.text_input("8. What conclusions can be drawn?")
            q9 = st.text_input("9. What future research could be conducted?")
        submitted = st.form_submit_button("Generate Article")
    if submitted:
        if language == "Russian":
            article_text = (
                f"Сгенерируйте подробную и комплексную научную статью в академическом стиле с гуманизированным тоном, "
                f"используя структуру IMRAD. Текст должен быть сплошным, без использования списков. "
                f"Представьте введение, описание методов, результаты, обсуждение и заключение. "
                f"В статье включите цитирования в квадратных скобках и литературный обзор с 10 источниками. "
                f"Введение: {q1} {q2} {q3}. Методы: {q4} {q5}. Результаты: {q6}. Обсуждение: {q7}. "
                f"Заключение: {q8} {q9}"
            )
        else:
            article_text = (
                f"Generate a detailed and comprehensive scientific article in academic style with a humanized tone, "
                f"using the IMRAD structure. The text should be continuous without bullet lists. "
                f"Present an introduction, methods description, results, discussion, and conclusion. "
                f"Include citations in square brackets and a literature overview with 10 references. "
                f"Introduction: {q1} {q2} {q3}. Methods: {q4} {q5}. Results: {q6}. Discussion: {q7}. "
                f"Conclusion: {q8} {q9}"
            )
        references = search_literature(q3, rows=10)

elif input_mode == "Voice Input (Real-Time)":
    st.header("Answer the following questions verbally:")
    st.write("Please answer the following questions one by one using your microphone:")
    if language == "Russian":
        st.write("1. Какую проблему вы исследуете?")
        st.write("2. Почему это важно?")
        st.write("3. Какие исследования уже проводились по данной теме?")
        st.write("4. Какие методы вы использовали?")
        st.write("5. Как вы собирали данные?")
        st.write("6. Какие результаты вы получили?")
        st.write("7. Соответствуют ли они ожиданиям или другим исследованиям?")
        st.write("8. Какие выводы можно сделать?")
        st.write("9. Какие дальнейшие исследования можно провести?")
    else:
        st.write("1. What problem are you researching?")
        st.write("2. Why is it important?")
        st.write("3. What previous research exists on this topic?")
        st.write("4. What methods did you use?")
        st.write("5. How did you collect data?")
        st.write("6. What results did you obtain?")
        st.write("7. Do they match expectations or other studies?")
        st.write("8. What conclusions can be drawn?")
        st.write("9. What future research could be conducted?")
    
    st.info("Click the button below to start real-time voice recording.")
    webrtc_ctx = webrtc_streamer(key="voice", 
                                 audio_processor_factory=VoskAudioProcessor,
                                 media_stream_constraints={"audio": True, "video": False})
    if webrtc_ctx.audio_processor:
        if st.button("Stop Recording and Use Transcript"):
            article_text = webrtc_ctx.audio_processor.transcript
            article_text = st.text_area("Edit transcript if needed:", article_text, height=200)
            submitted = True
            if not article_text or "research" not in article_text.lower():
                q3_voice = st.text_input("Enter your previous research query for literature search:")
                if q3_voice:
                    references = search_literature(q3_voice, rows=10)
            else:
                references = search_literature(article_text, rows=10)

if submitted and article_text:
    with st.spinner("Generating article..."):
        generated_article = generate_text(article_text)
    st.subheader("Generated Article")
    st.write(generated_article)
    
    vis_data = extract_visualization_data(generated_article)
    if vis_data:
        bar_chart_path, line_chart_path = plot_generated_data(vis_data)
    else:
        st.info("No visualization data extracted. Using sample data.")
        sample_data = [
            {"Category": "A", "Value": 10},
            {"Category": "B", "Value": 20},
            {"Category": "C", "Value": 15},
            {"Category": "D", "Value": 30}
        ]
        vis_data = sample_data
        bar_chart_path, line_chart_path = plot_generated_data(sample_data)
    
    if references:
        st.header("Literature Overview")
        for ref in references:
            st.markdown(f"**[{ref['index']}] {ref['title']}**  \nJournal: {ref['journal']}  \nDOI: {ref['doi']}  \nOverview: {ref['overview']}\n")
    
    st.header("Export Article")
    docx_file = export_document(generated_article, vis_data, references)
    if docx_file and os.path.exists(docx_file):
        with open(docx_file, "rb") as f:
            st.download_button("Download DOCX", f, file_name=docx_file, key="docx_download")
    
    chat_entry = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "username": st.session_state.username,
        "input": article_text,
        "generated": generated_article
    }
    save_chat_history(chat_entry)
